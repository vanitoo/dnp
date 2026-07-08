from __future__ import annotations

import re
import sys
import traceback
from dataclasses import dataclass
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Any

import openpyxl
from docx import Document

MONTH_NAMES = {
    1: "январь", 2: "февраль", 3: "март", 4: "апрель", 5: "май", 6: "июнь",
    7: "июль", 8: "август", 9: "сентябрь", 10: "октябрь", 11: "ноябрь", 12: "декабрь",
}

DEFAULT_TARIFFS = {"Т1": 9.67, "Т2": 3.51, "Т3": 6.77}
DEFAULT_WATER_TARIFF = 53.0

@dataclass
class ReceiptRow:
    plot: str
    tariff_name: str
    prev_value: float | None
    curr_value: float | None
    consumption: float | None
    tariff: float
    amount: float | None

@dataclass
class PlotReceipt:
    plot: str
    month: int
    year: int
    rows: list[ReceiptRow]
    water_prev: float | None = None
    water_curr: float | None = None
    water_consumption: float | None = None
    water_tariff: float = DEFAULT_WATER_TARIFF
    water_amount: float | None = None

    @property
    def total_electricity(self) -> float:
        return money_round(sum(r.amount or 0 for r in self.rows))

    @property
    def total(self) -> float:
        return money_round(self.total_electricity + (self.water_amount or 0))


def safe_num(value: Any) -> float | None:
    if value is None or value == "":
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).replace("\xa0", "").replace(" ", "").replace(",", ".")
    try:
        return float(text)
    except ValueError:
        return None


def money_round(value: float | int | Decimal) -> float:
    return float(Decimal(str(value)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP))


def fmt_num(value: float | None, digits: int = 0) -> str:
    if value is None:
        return ""
    if digits == 0 and abs(value - round(value)) < 0.00001:
        return str(int(round(value)))
    return f"{value:.{digits}f}".replace(".", ",")


def fmt_money(value: float | None) -> str:
    if value is None:
        return ""
    return f"{money_round(value):.2f}".replace(".", ",")


def normalize_plot(value: Any) -> str | None:
    if value is None or value == "":
        return None
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).replace("\xa0", "").strip()


def choose_sheet(wb: openpyxl.Workbook):
    best_ws = None
    best_count = -1
    for ws in wb.worksheets:
        count = 0
        for row in range(1, ws.max_row + 1):
            if normalize_plot(ws.cell(row, 1).value):
                count += 1
        if count > best_count:
            best_ws = ws
            best_count = count
    return best_ws


def month_columns(header_row: int, ws) -> dict[int, int]:
    result = {}
    for col in range(1, ws.max_column + 1):
        num = safe_num(ws.cell(header_row, col).value)
        if num is not None and int(num) == num and 1 <= int(num) <= 12:
            result[int(num)] = col
    return result


def find_tariff_row(ws, start_row: int, tariff_name: str) -> int | None:
    target = tariff_name.upper().replace(" ", "")
    for row in range(start_row + 1, min(start_row + 8, ws.max_row + 1)):
        value = ws.cell(row, 2).value
        if value is None:
            continue
        norm = str(value).upper().replace(" ", "")
        if norm == target:
            return row
    return None


def parse_tariffs_from_template(template_path: str | Path) -> tuple[dict[str, float], float]:
    """Берет тарифы из DOCX как значения по умолчанию для GUI."""
    tariffs = dict(DEFAULT_TARIFFS)
    water_tariff = DEFAULT_WATER_TARIFF
    try:
        doc = Document(template_path)
        for table in doc.tables:
            text = "\n".join(cell.text for row in table.rows for cell in row.cells)
            if "Электроэнергия" in text:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if len(cells) >= 5:
                        name = cells[1].replace(" ", "").upper()
                        val = safe_num(cells[4])
                        if name in tariffs and val is not None:
                            tariffs[name] = val
            if "Водоотведение" in text:
                nums = re.findall(r"\d+(?:[,.]\d+)?", text)
                if nums:
                    water_tariff = float(nums[-1].replace(",", "."))
    except Exception:
        pass
    return tariffs, water_tariff


def load_receipts_from_excel(
    xlsx_path: str | Path,
    month: int,
    year: int,
    tariffs: dict[str, float] | None = None,
    water_tariff: float = DEFAULT_WATER_TARIFF,
    water_consumption: float = 0.0,
) -> list[PlotReceipt]:
    tariffs = tariffs or dict(DEFAULT_TARIFFS)
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = choose_sheet(wb)
    if ws is None:
        raise RuntimeError("В Excel не найден лист с данными")

    receipts: list[PlotReceipt] = []
    previous_month = 12 if month == 1 else month - 1

    for row in range(1, ws.max_row + 1):
        plot = normalize_plot(ws.cell(row, 1).value)
        if not plot:
            continue
        cols = month_columns(row, ws)
        if month not in cols:
            continue
        curr_col = cols[month]
        prev_col = cols.get(previous_month)
        if prev_col is None:
            raise RuntimeError(f"Для участка {plot} нет колонки предыдущего месяца {previous_month}.")

        rows: list[ReceiptRow] = []
        for tariff_name in ["Т1", "Т2", "Т3"]:
            tariff_value = float(tariffs.get(tariff_name, DEFAULT_TARIFFS[tariff_name]))
            tariff_row = find_tariff_row(ws, row, tariff_name)
            if tariff_row is None:
                prev_value = curr_value = consumption = amount = None
            else:
                prev_value = safe_num(ws.cell(tariff_row, prev_col).value)
                curr_value = safe_num(ws.cell(tariff_row, curr_col).value)
                if prev_value is None or curr_value is None:
                    consumption = amount = None
                else:
                    consumption = curr_value - prev_value
                    amount = money_round(consumption * tariff_value)
            rows.append(ReceiptRow(plot, tariff_name, prev_value, curr_value, consumption, tariff_value, amount))

        wc = float(water_consumption or 0)
        water_amount = money_round(wc * water_tariff) if wc else 0.0
        receipts.append(PlotReceipt(plot, month, year, rows, None, None, wc if wc else None, water_tariff, water_amount))
    return receipts


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.add_run(text)


def amount_to_words_ru(amount: float) -> str:
    amount_dec = Decimal(str(amount)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    rub = int(amount_dec)
    kop = int((amount_dec - rub) * 100)
    return f"{number_to_words_ru(rub, female=False).capitalize()} {plural_ru(rub, 'рубль', 'рубля', 'рублей')} {kop:02d} {plural_ru(kop, 'копейка', 'копейки', 'копеек')}"


def plural_ru(n: int, one: str, two: str, five: str) -> str:
    n = abs(n) % 100
    if 11 <= n <= 19:
        return five
    n = n % 10
    if n == 1:
        return one
    if 2 <= n <= 4:
        return two
    return five


def number_to_words_ru(n: int, female: bool = False) -> str:
    if n == 0:
        return "ноль"
    ones_m = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    ones_f = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    teens = ["десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать", "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
    tens = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят", "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
    hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот", "шестьсот", "семьсот", "восемьсот", "девятьсот"]

    def triad_to_words(x: int, fem: bool) -> list[str]:
        out: list[str] = []
        out.append(hundreds[x // 100])
        last = x % 100
        if 10 <= last <= 19:
            out.append(teens[last - 10])
        else:
            out.append(tens[last // 10])
            out.append((ones_f if fem else ones_m)[last % 10])
        return [w for w in out if w]

    groups = [
        (1_000_000_000, False, ("миллиард", "миллиарда", "миллиардов")),
        (1_000_000, False, ("миллион", "миллиона", "миллионов")),
        (1_000, True, ("тысяча", "тысячи", "тысяч")),
        (1, female, ("", "", "")),
    ]
    words: list[str] = []
    rest = n
    for div, fem, forms in groups:
        part = rest // div
        rest %= div
        if part == 0:
            continue
        words.extend(triad_to_words(part, fem))
        if div != 1:
            words.append(plural_ru(part, *forms))
    return " ".join(words)


def replace_paragraph_text(doc: Document, receipt: PlotReceipt) -> None:
    month_name = MONTH_NAMES[receipt.month]
    for p in doc.paragraphs:
        text = p.text
        if "Участок №" in text:
            p.text = f"                      Участок № {receipt.plot}    « {month_name} » {receipt.year} год"
        elif "Сумма оплаты:" in text:
            p.text = f"                      Сумма оплаты: {amount_to_words_ru(receipt.total)}"


def fill_electricity_table(doc: Document, receipt: PlotReceipt) -> None:
    table = None
    for t in doc.tables:
        if len(t.rows) >= 6 and len(t.columns) >= 6 and "Электроэнергия" in t.cell(1, 0).text:
            table = t
            break
    if table is None:
        raise RuntimeError("В шаблоне не найдена таблица 'Электроэнергия'")
    data_rows = [2, 3, 4]
    for idx, rr in enumerate(receipt.rows):
        r = data_rows[idx]
        set_cell_text(table.cell(r, 1), fmt_num(rr.curr_value))
        set_cell_text(table.cell(r, 2), fmt_num(rr.prev_value))
        set_cell_text(table.cell(r, 3), fmt_num(rr.consumption))
        set_cell_text(table.cell(r, 4), fmt_money(rr.tariff))
        set_cell_text(table.cell(r, 5), fmt_money(rr.amount))
    set_cell_text(table.cell(5, 5), f"{fmt_money(receipt.total_electricity)} рублей")


def fill_water_table(doc: Document, receipt: PlotReceipt) -> None:
    table = None
    for t in doc.tables:
        if len(t.rows) >= 4 and len(t.columns) >= 6 and "Водоотведение" in t.cell(1, 0).text:
            table = t
            break
    if table is None:
        return
    # Если водоотведение пока не ведется в Excel — можно задать фиксированный расход в GUI.
    set_cell_text(table.cell(2, 1), fmt_num(receipt.water_curr))
    set_cell_text(table.cell(2, 2), fmt_num(receipt.water_prev))
    set_cell_text(table.cell(2, 3), fmt_num(receipt.water_consumption))
    set_cell_text(table.cell(2, 4), f"{fmt_money(receipt.water_tariff)} рублей")
    set_cell_text(table.cell(2, 5), fmt_money(receipt.water_amount))
    set_cell_text(table.cell(3, 5), f"{fmt_money(receipt.water_amount)} рублей")


def sanitize_filename(value: str) -> str:
    return re.sub(r"[^0-9A-Za-zА-Яа-я_. -]+", "_", value).strip()




def main_cli(argv: list[str]) -> int:
    import argparse
    parser = argparse.ArgumentParser(description="Генератор квитанций ДНП из Excel")
    parser.add_argument("--xlsx", required=True, help="Путь к Excel-файлу")
    parser.add_argument("--template", required=True, help="Путь к DOCX-шаблону квитанции")
    parser.add_argument("--month", type=int, required=True, help="Месяц числом: 1-12")
    parser.add_argument("--year", type=int, required=True, help="Год, например 2025")
    parser.add_argument("--out", required=True, help="Папка для готовых квитанций")
    parser.add_argument("--t1", type=float, default=None)
    parser.add_argument("--t2", type=float, default=None)
    parser.add_argument("--t3", type=float, default=None)
    parser.add_argument("--water-tariff", type=float, default=None)
    parser.add_argument("--water-consumption", type=float, default=0.0)
    parser.add_argument("--expected-count", type=int, default=39, help="Ожидаемое число квитанций для контроля")
    args = parser.parse_args(argv)
    tpl_tariffs, tpl_water = parse_tariffs_from_template(args.template)
    tariffs = {"Т1": args.t1 if args.t1 is not None else tpl_tariffs["Т1"],
               "Т2": args.t2 if args.t2 is not None else tpl_tariffs["Т2"],
               "Т3": args.t3 if args.t3 is not None else tpl_tariffs["Т3"]}
    water_tariff = args.water_tariff if args.water_tariff is not None else tpl_water
    paths = generate_all(args.xlsx, args.template, args.month, args.year, args.out, tariffs, water_tariff, args.water_consumption)
    print(f"Готово: создано файлов: {len(paths)}")
    if args.expected_count and len(paths) != args.expected_count:
        print(f"Внимание: ожидалось {args.expected_count}, а создано {len(paths)}. Проверьте, все ли участки есть в Excel.")
    for p in paths:
        print(p)
    return 0



MARKER_ALIASES = {
    "{{P}}": "plot", "{{LOT}}": "plot", "{{M}}": "month_name", "{{MONTH}}": "month_name",
    "{{Y}}": "year", "{{YEAR}}": "year", "{{SW}}": "total_words", "{{TOTAL_TEXT}}": "total_words",
    "{{ES}}": "electricity_total", "{{WT}}": "water_total", "{{TOTAL}}": "total",
    "{{T1C}}": "t1_curr", "{{T1P}}": "t1_prev", "{{T1R}}": "t1_usage", "{{R1}}": "t1_rate", "{{S1}}": "t1_sum",
    "{{T2C}}": "t2_curr", "{{T2P}}": "t2_prev", "{{T2R}}": "t2_usage", "{{R2}}": "t2_rate", "{{S2}}": "t2_sum",
    "{{T3C}}": "t3_curr", "{{T3P}}": "t3_prev", "{{T3R}}": "t3_usage", "{{R3}}": "t3_rate", "{{S3}}": "t3_sum",
    "{{WC}}": "water_curr", "{{WP}}": "water_prev", "{{WR}}": "water_usage", "{{RW}}": "water_rate", "{{WS}}": "water_sum",
}


def replace_text_keep_first_run(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = text


def docx_has_markers(doc: Document) -> bool:
    pat = re.compile(r"\{\{[A-Za-z0-9_]+\}\}")
    if any(pat.search(p.text) for p in doc.paragraphs):
        return True
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if pat.search(cell.text):
                    return True
    return False


def receipt_marker_values(receipt: PlotReceipt) -> dict[str, str]:
    rows = {r.tariff_name: r for r in receipt.rows}
    def row(name: str) -> ReceiptRow:
        return rows.get(name) or ReceiptRow(receipt.plot, name, None, None, None, 0, None)
    t1, t2, t3 = row("Т1"), row("Т2"), row("Т3")
    values = {
        "plot": receipt.plot, "month_name": MONTH_NAMES[receipt.month], "year": str(receipt.year),
        "total_words": amount_to_words_ru(receipt.total), "electricity_total": fmt_money(receipt.total_electricity),
        "water_total": fmt_money(receipt.water_amount), "total": fmt_money(receipt.total),
        "t1_curr": fmt_num(t1.curr_value), "t1_prev": fmt_num(t1.prev_value), "t1_usage": fmt_num(t1.consumption), "t1_rate": fmt_money(t1.tariff), "t1_sum": fmt_money(t1.amount),
        "t2_curr": fmt_num(t2.curr_value), "t2_prev": fmt_num(t2.prev_value), "t2_usage": fmt_num(t2.consumption), "t2_rate": fmt_money(t2.tariff), "t2_sum": fmt_money(t2.amount),
        "t3_curr": fmt_num(t3.curr_value), "t3_prev": fmt_num(t3.prev_value), "t3_usage": fmt_num(t3.consumption), "t3_rate": fmt_money(t3.tariff), "t3_sum": fmt_money(t3.amount),
        "water_curr": fmt_num(receipt.water_curr), "water_prev": fmt_num(receipt.water_prev), "water_usage": fmt_num(receipt.water_consumption), "water_rate": fmt_money(receipt.water_tariff), "water_sum": fmt_money(receipt.water_amount),
    }
    return {marker: values[key] for marker, key in MARKER_ALIASES.items()}


def replace_markers_in_paragraph(paragraph, values: dict[str, str]) -> None:
    text = paragraph.text
    new = text
    for marker, val in values.items():
        new = new.replace(marker, val)
    if new != text:
        replace_text_keep_first_run(paragraph, new)


def fill_marker_template(doc: Document, receipt: PlotReceipt) -> None:
    values = receipt_marker_values(receipt)
    for p in doc.paragraphs:
        replace_markers_in_paragraph(p, values)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_markers_in_paragraph(p, values)


def make_receipt_docx(template_path: str | Path, receipt: PlotReceipt, output_dir: str | Path) -> Path:
    doc = Document(template_path)
    if docx_has_markers(doc):
        fill_marker_template(doc, receipt)
    else:
        replace_paragraph_text(doc, receipt)
        fill_electricity_table(doc, receipt)
        fill_water_table(doc, receipt)
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"Участок_{sanitize_filename(receipt.plot)}_квитанция_{receipt.year}_{receipt.month:02d}.docx"
    doc.save(out_path)
    return out_path


def set_cell_text_preserve(cell, text: str) -> None:
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def make_template_with_markers(input_docx: str | Path, output_docx: str | Path) -> Path:
    doc = Document(input_docx)
    for p in doc.paragraphs:
        if "Участок №" in p.text:
            replace_text_keep_first_run(p, "                      Участок № {{P}}    {{M}} {{Y}} год")
        elif "Сумма оплаты:" in p.text:
            replace_text_keep_first_run(p, "                      Сумма оплаты: {{SW}}")
    for t in doc.tables:
        if len(t.rows) >= 6 and len(t.columns) >= 6 and "Электроэнергия" in t.cell(1, 0).text:
            data = {(2,1):"{{T1C}}",(2,2):"{{T1P}}",(2,3):"{{T1R}}",(2,4):"{{R1}}",(2,5):"{{S1}}",(3,1):"{{T2C}}",(3,2):"{{T2P}}",(3,3):"{{T2R}}",(3,4):"{{R2}}",(3,5):"{{S2}}",(4,1):"{{T3C}}",(4,2):"{{T3P}}",(4,3):"{{T3R}}",(4,4):"{{R3}}",(4,5):"{{S3}}",(5,5):"{{ES}} руб."}
            for (r,c), v in data.items():
                set_cell_text_preserve(t.cell(r,c), v)
        if len(t.rows) >= 4 and len(t.columns) >= 6 and "Водоотведение" in t.cell(1, 0).text:
            data = {(2,1):"{{WC}}",(2,2):"{{WP}}",(2,3):"{{WR}}",(2,4):"{{RW}}",(2,5):"{{WS}}",(3,5):"{{WT}} руб."}
            for (r,c), v in data.items():
                set_cell_text_preserve(t.cell(r,c), v)
    output_docx = Path(output_docx)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    return output_docx


def convert_docx_to_pdf(docx_path: str | Path, pdf_path: str | Path | None = None) -> Path:
    import shutil, subprocess
    docx_path = Path(docx_path).resolve()
    pdf_path = Path(pdf_path).resolve() if pdf_path else docx_path.with_suffix('.pdf')
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    if sys.platform.startswith('win'):
        try:
            import win32com.client  # type: ignore
            word = win32com.client.DispatchEx('Word.Application')
            word.Visible = False
            doc = word.Documents.Open(str(docx_path))
            doc.SaveAs(str(pdf_path), FileFormat=17)
            doc.Close(False)
            word.Quit()
            return pdf_path
        except Exception:
            pass
    soffice = shutil.which('soffice') or shutil.which('libreoffice')
    if not soffice:
        raise RuntimeError('PDF не создан: не найден Microsoft Word или LibreOffice.')
    subprocess.run([soffice, '--headless', '--convert-to', 'pdf', '--outdir', str(pdf_path.parent), str(docx_path)], check=True)
    generated = pdf_path.parent / docx_path.with_suffix('.pdf').name
    if generated != pdf_path and generated.exists():
        generated.replace(pdf_path)
    return pdf_path


def generate_all(xlsx_path: str | Path, template_path: str | Path, month: int, year: int, output_dir: str | Path,
                 tariffs: dict[str, float] | None = None, water_tariff: float = DEFAULT_WATER_TARIFF,
                 water_consumption: float = 0.0, save_pdf: bool = False) -> list[Path]:
    receipts = load_receipts_from_excel(xlsx_path, month, year, tariffs, water_tariff, water_consumption)
    if not receipts:
        raise RuntimeError("Не найдено ни одного участка для выбранного месяца")
    paths: list[Path] = []
    for r in receipts:
        docx = make_receipt_docx(template_path, r, output_dir)
        paths.append(docx)
        if save_pdf:
            paths.append(convert_docx_to_pdf(docx))
    return paths

EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")

def load_email_recipients_from_excel(xlsx_path: str | Path, use_test_emails: bool = True) -> dict[str, str]:
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = choose_sheet(wb)
    if ws is None:
        raise RuntimeError("В Excel не найден лист с данными")
    result: dict[str, str] = {}
    for row in range(1, ws.max_row + 1):
        plot = normalize_plot(ws.cell(row, 1).value)
        if not plot or not month_columns(row, ws):
            continue
        email = ""
        for r in range(row, min(row + 8, ws.max_row + 1)):
            for c in range(1, ws.max_column + 1):
                m = EMAIL_RE.search(str(ws.cell(r, c).value or ""))
                if m:
                    email = m.group(0); break
            if email: break
        if not email and use_test_emails:
            safe = re.sub(r"\D+", "", plot) or sanitize_filename(plot).replace('_', '') or '0'
            email = f"{safe}@mail.ru"
        result[plot] = email
    return result


def extract_plot_from_receipt_filename(path: str | Path) -> str | None:
    m = re.search(r"Участок_(.+?)_квитанция_", Path(path).stem)
    return m.group(1).replace('_', '/').strip() if m else None


def send_email_with_attachment(smtp_settings: dict[str, Any], to_addr: str, subject: str, body: str, attachment_path: str | Path) -> None:
    import smtplib, ssl
    from email.message import EmailMessage
    path = Path(attachment_path)
    msg = EmailMessage()
    msg['From'] = smtp_settings['from_email']
    msg['To'] = to_addr
    msg['Subject'] = subject
    msg.set_content(body)
    subtype = 'pdf' if path.suffix.lower() == '.pdf' else 'vnd.openxmlformats-officedocument.wordprocessingml.document'
    msg.add_attachment(path.read_bytes(), maintype='application', subtype=subtype, filename=path.name)
    host = smtp_settings['host']; port = int(smtp_settings['port'])
    username = smtp_settings.get('username') or smtp_settings.get('from_email')
    password = smtp_settings.get('password') or ''
    security = smtp_settings.get('security', 'ssl')
    if security == 'ssl':
        with smtplib.SMTP_SSL(host, port, timeout=30, context=ssl.create_default_context()) as s:
            if username: s.login(username, password)
            s.send_message(msg)
    else:
        with smtplib.SMTP(host, port, timeout=30) as s:
            s.ehlo()
            if security == 'starttls':
                s.starttls(context=ssl.create_default_context()); s.ehlo()
            if username: s.login(username, password)
            s.send_message(msg)


def main_gui() -> None:
    import json, threading
    import tkinter as tk
    from tkinter import filedialog, messagebox, ttk

    CONFIG_PATH = Path.home() / '.dnp_receipts_settings.json'

    def load_config() -> dict[str, Any]:
        try:
            return json.loads(CONFIG_PATH.read_text(encoding='utf-8')) if CONFIG_PATH.exists() else {}
        except Exception:
            return {}

    def save_config(data: dict[str, Any]) -> None:
        CONFIG_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')

    cfg = load_config()
    root = tk.Tk()
    root.title('ДНП Комфорт — квитанции')
    root.geometry('980x680')
    nb = ttk.Notebook(root); nb.pack(fill='both', expand=True, padx=8, pady=8)
    tab_gen = ttk.Frame(nb); tab_mail = ttk.Frame(nb); tab_smtp = ttk.Frame(nb)
    nb.add(tab_gen, text='Генерация'); nb.add(tab_mail, text='Отправка на почту'); nb.add(tab_smtp, text='Настройки почты')

    vars_ = {
        'xlsx': tk.StringVar(value=cfg.get('xlsx','')),
        'template': tk.StringVar(value=cfg.get('template','')),
        'out': tk.StringVar(value=cfg.get('out', str(Path.cwd() / 'Готовые квитанции'))),
        'month': tk.IntVar(value=int(cfg.get('month', 12))),
        'year': tk.IntVar(value=int(cfg.get('year', 2025))),
        't1': tk.StringVar(value=cfg.get('t1', '9,67')),
        't2': tk.StringVar(value=cfg.get('t2', '3,51')),
        't3': tk.StringVar(value=cfg.get('t3', '6,77')),
        'water_tariff': tk.StringVar(value=cfg.get('water_tariff', '53,00')),
        'water_consumption': tk.StringVar(value=cfg.get('water_consumption', '0')),
        'save_pdf': tk.BooleanVar(value=bool(cfg.get('save_pdf', False))),
        'mail_folder': tk.StringVar(value=cfg.get('mail_folder', str(Path.cwd() / 'Готовые квитанции'))),
        'mail_use_test': tk.BooleanVar(value=bool(cfg.get('mail_use_test', True))),
    }
    smtp_vars = {
        'host': tk.StringVar(value=cfg.get('smtp', {}).get('host', 'smtp.mail.ru')),
        'port': tk.StringVar(value=str(cfg.get('smtp', {}).get('port', '465'))),
        'security': tk.StringVar(value=cfg.get('smtp', {}).get('security', 'ssl')),
        'from_email': tk.StringVar(value=cfg.get('smtp', {}).get('from_email', '')),
        'username': tk.StringVar(value=cfg.get('smtp', {}).get('username', '')),
        'password': tk.StringVar(value=cfg.get('smtp', {}).get('password', '')),
        'subject': tk.StringVar(value=cfg.get('smtp', {}).get('subject', 'Квитанция ДНП «Комфорт» за {month} {year}')),
    }

    def persist() -> None:
        data = {k: v.get() for k, v in vars_.items()}
        data['smtp'] = {k: v.get() for k, v in smtp_vars.items()}
        save_config(data)

    def parse_float(name: str) -> float:
        return float((vars_[name].get() or '0').replace(' ', '').replace(',', '.'))

    def add_file_row(parent, label, var, row_num, kind):
        ttk.Label(parent, text=label).grid(row=row_num, column=0, sticky='w', padx=10, pady=6)
        ttk.Entry(parent, textvariable=var, width=88).grid(row=row_num, column=1, sticky='we', padx=6, pady=6)
        def browse():
            if kind == 'xlsx':
                path = filedialog.askopenfilename(filetypes=[('Excel', '*.xlsx')])
            elif kind == 'docx':
                path = filedialog.askopenfilename(filetypes=[('Word', '*.docx')])
            else:
                path = filedialog.askdirectory()
            if path:
                var.set(path)
                if kind == 'docx':
                    tariffs, water = parse_tariffs_from_template(path)
                    vars_['t1'].set(fmt_money(tariffs['Т1']))
                    vars_['t2'].set(fmt_money(tariffs['Т2']))
                    vars_['t3'].set(fmt_money(tariffs['Т3']))
                    vars_['water_tariff'].set(fmt_money(water))
                persist()
        ttk.Button(parent, text='Выбрать', command=browse).grid(row=row_num, column=2, padx=6, pady=6)
        parent.columnconfigure(1, weight=1)

    add_file_row(tab_gen, 'Excel с показаниями', vars_['xlsx'], 0, 'xlsx')
    add_file_row(tab_gen, 'Шаблон квитанции DOCX', vars_['template'], 1, 'docx')
    add_file_row(tab_gen, 'Папка вывода', vars_['out'], 2, 'folder')
    ttk.Label(tab_gen, text='Месяц').grid(row=3, column=0, sticky='w', padx=10, pady=6)
    ttk.Spinbox(tab_gen, from_=1, to=12, textvariable=vars_['month'], width=8).grid(row=3, column=1, sticky='w', padx=6, pady=6)
    ttk.Label(tab_gen, text='Год').grid(row=4, column=0, sticky='w', padx=10, pady=6)
    ttk.Spinbox(tab_gen, from_=2020, to=2035, textvariable=vars_['year'], width=8).grid(row=4, column=1, sticky='w', padx=6, pady=6)
    tf = ttk.LabelFrame(tab_gen, text='Тарифы'); tf.grid(row=5, column=0, columnspan=3, sticky='we', padx=10, pady=10)
    for i, key in enumerate(['t1','t2','t3']):
        ttk.Label(tf, text=key.upper()).grid(row=0, column=i*2, padx=8, pady=8)
        ttk.Entry(tf, textvariable=vars_[key], width=12).grid(row=0, column=i*2+1, padx=4, pady=8)
    ttk.Label(tf, text='Водоотведение тариф').grid(row=0, column=6, padx=8, pady=8)
    ttk.Entry(tf, textvariable=vars_['water_tariff'], width=12).grid(row=0, column=7, padx=4, pady=8)
    ttk.Label(tf, text='Расход м3').grid(row=0, column=8, padx=8, pady=8)
    ttk.Entry(tf, textvariable=vars_['water_consumption'], width=12).grid(row=0, column=9, padx=4, pady=8)
    ttk.Checkbutton(tab_gen, text='Сохранять PDF дополнительно к DOCX', variable=vars_['save_pdf'], command=persist).grid(row=6, column=1, sticky='w', padx=6, pady=6)
    gen_status = tk.StringVar(value='PDF по умолчанию выключен. Для PDF нужен Microsoft Word или LibreOffice.')
    ttk.Label(tab_gen, textvariable=gen_status, wraplength=850).grid(row=9, column=0, columnspan=3, sticky='we', padx=10, pady=8)

    def create_marked_template():
        try:
            if not vars_['template'].get():
                raise RuntimeError('Сначала выберите исходный DOCX-шаблон')
            dst = filedialog.asksaveasfilename(defaultextension='.docx', initialfile='Квитанция_шаблон_с_метками.docx', filetypes=[('Word', '*.docx')])
            if not dst: return
            out = make_template_with_markers(vars_['template'].get(), dst)
            vars_['template'].set(str(out)); persist()
            messagebox.showinfo('Готово', f'Шаблон с метками создан:\n{out}')
        except Exception as e:
            traceback.print_exc(); messagebox.showerror('Ошибка', str(e))

    def run_generation():
        try:
            persist()
            tariffs = {'Т1': parse_float('t1'), 'Т2': parse_float('t2'), 'Т3': parse_float('t3')}
            paths = generate_all(vars_['xlsx'].get(), vars_['template'].get(), int(vars_['month'].get()), int(vars_['year'].get()), vars_['out'].get(), tariffs=tariffs, water_tariff=parse_float('water_tariff'), water_consumption=parse_float('water_consumption'), save_pdf=vars_['save_pdf'].get())
            docx_count = len([p for p in paths if p.suffix.lower()=='.docx']); pdf_count = len([p for p in paths if p.suffix.lower()=='.pdf'])
            gen_status.set(f'Готово. DOCX: {docx_count}, PDF: {pdf_count}. Папка: {vars_["out"].get()}')
            msg = f'Создано DOCX: {docx_count}\nСоздано PDF: {pdf_count}'
            if docx_count != 39: msg += '\n\nВнимание: ожидалось 39 участков. Проверьте Excel.'
            messagebox.showinfo('Готово', msg)
        except Exception as e:
            traceback.print_exc(); gen_status.set(f'Ошибка: {e}'); messagebox.showerror('Ошибка', str(e))

    bf = ttk.Frame(tab_gen); bf.grid(row=7, column=1, sticky='we', padx=6, pady=16)
    ttk.Button(bf, text='Создать шаблон с метками', command=create_marked_template).pack(side='left', padx=4)
    ttk.Button(bf, text='Сформировать квитанции', command=run_generation).pack(side='left', padx=4)

    # SMTP settings
    rows = [('SMTP сервер','host'),('Порт','port'),('Email отправителя','from_email'),('Логин','username'),('Пароль приложения','password'),('Тема письма','subject')]
    for i, (label, key) in enumerate(rows):
        ttk.Label(tab_smtp, text=label).grid(row=i, column=0, sticky='w', padx=10, pady=6)
        ttk.Entry(tab_smtp, textvariable=smtp_vars[key], width=80, show='*' if key == 'password' else None).grid(row=i, column=1, sticky='we', padx=6, pady=6)
    ttk.Label(tab_smtp, text='Шифрование').grid(row=len(rows), column=0, sticky='w', padx=10, pady=6)
    ttk.Combobox(tab_smtp, textvariable=smtp_vars['security'], values=['ssl','starttls','none'], width=12, state='readonly').grid(row=len(rows), column=1, sticky='w', padx=6, pady=6)
    tab_smtp.columnconfigure(1, weight=1)
    ttk.Label(tab_smtp, text='Важно: для Mail.ru/Яндекс/Gmail обычно нужен пароль приложения, а не обычный пароль.', wraplength=820).grid(row=len(rows)+1, column=0, columnspan=2, sticky='we', padx=10, pady=12)
    ttk.Button(tab_smtp, text='Сохранить настройки', command=lambda: (persist(), messagebox.showinfo('Готово','Настройки сохранены'))).grid(row=len(rows)+2, column=1, sticky='w', padx=6, pady=8)

    # Mail tab
    add_file_row(tab_mail, 'Excel с участками/email', vars_['xlsx'], 0, 'xlsx')
    add_file_row(tab_mail, 'Папка с готовыми квитанциями', vars_['mail_folder'], 1, 'folder')
    ttk.Checkbutton(tab_mail, text='Если email не найден — использовать тестовый адрес вида 1@mail.ru', variable=vars_['mail_use_test'], command=persist).grid(row=2, column=1, sticky='w', padx=6, pady=6)
    mail_status = tk.StringVar(value='Сначала нажмите «Проверить список». Отправка идет последовательно: одно письмо за другим.')
    ttk.Label(tab_mail, textvariable=mail_status, wraplength=850).grid(row=6, column=0, columnspan=3, sticky='we', padx=10, pady=8)
    tree = ttk.Treeview(tab_mail, columns=('plot','email','file','status'), show='headings', height=18)
    for col, title, w in [('plot','Участок',90),('email','Email',220),('file','Файл',430),('status','Статус',170)]:
        tree.heading(col, text=title); tree.column(col, width=w)
    tree.grid(row=5, column=0, columnspan=3, sticky='nsew', padx=10, pady=8)
    tab_mail.rowconfigure(5, weight=1); tab_mail.columnconfigure(1, weight=1)
    send_items: list[tuple[str, str, Path]] = []

    def build_send_list() -> list[tuple[str, str, Path]]:
        recipients = load_email_recipients_from_excel(vars_['xlsx'].get(), vars_['mail_use_test'].get())
        folder = Path(vars_['mail_folder'].get())
        if not folder.exists(): raise RuntimeError('Папка с квитанциями не найдена')
        files = sorted(list(folder.glob('*.pdf')) or list(folder.glob('*.docx')))
        result = []
        for f in files:
            plot = extract_plot_from_receipt_filename(f)
            if not plot: continue
            email = recipients.get(plot, '')
            if not email and vars_['mail_use_test'].get():
                email = f"{re.sub(r'[^0-9]+', '', plot) or sanitize_filename(plot)}@mail.ru"
            result.append((plot, email, f))
        return result

    def preview_send_list():
        nonlocal send_items
        try:
            persist(); tree.delete(*tree.get_children()); send_items = build_send_list()
            for plot, email, f in send_items:
                tree.insert('', 'end', values=(plot, email, f.name, 'готово' if email else 'нет email'))
            mail_status.set(f'Найдено файлов к отправке: {len(send_items)}')
        except Exception as e:
            traceback.print_exc(); messagebox.showerror('Ошибка', str(e)); mail_status.set(f'Ошибка: {e}')

    def smtp_settings() -> dict[str, Any]:
        persist(); return {k: v.get() for k, v in smtp_vars.items()}

    def send_all_thread():
        settings = smtp_settings(); month_name = MONTH_NAMES[int(vars_['month'].get())]; year = int(vars_['year'].get())
        ok = fail = 0; children = tree.get_children()
        for idx, item_id in enumerate(children):
            plot, email, filename, _ = tree.item(item_id, 'values')
            file_path = next((x[2] for x in send_items if x[2].name == filename), None)
            if not email or file_path is None:
                tree.set(item_id, 'status', 'пропущено'); fail += 1; continue
            try:
                tree.set(item_id, 'status', 'отправка...')
                subject = (settings.get('subject') or 'Квитанция').format(month=month_name, year=year, plot=plot)
                body = f'Здравствуйте.\n\nНаправляем квитанцию по участку № {plot} за {month_name} {year}.\n\nДНП «Комфорт»'
                send_email_with_attachment(settings, email, subject, body, file_path)
                tree.set(item_id, 'status', 'отправлено'); ok += 1
            except Exception as e:
                tree.set(item_id, 'status', f'ошибка: {str(e)[:80]}'); fail += 1
            mail_status.set(f'Отправка: {idx+1}/{len(children)}. Успешно: {ok}, ошибок: {fail}')
        mail_status.set(f'Отправка завершена. Успешно: {ok}, ошибок: {fail}')

    def send_all():
        if not send_items: preview_send_list()
        if not messagebox.askyesno('Подтверждение', 'Начать последовательную отправку квитанций?'):
            return
        threading.Thread(target=send_all_thread, daemon=True).start()

    mf = ttk.Frame(tab_mail); mf.grid(row=4, column=1, sticky='w', padx=6, pady=8)
    ttk.Button(mf, text='Проверить список', command=preview_send_list).pack(side='left', padx=4)
    ttk.Button(mf, text='Отправить последовательно', command=send_all).pack(side='left', padx=4)

    root.protocol('WM_DELETE_WINDOW', lambda: (persist(), root.destroy()))
    root.mainloop()


if __name__ == '__main__':
    if len(sys.argv) > 1:
        raise SystemExit(main_cli(sys.argv[1:]))
    main_gui()
